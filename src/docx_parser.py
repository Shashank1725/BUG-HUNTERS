"""
docx_parser.py
--------------
DOCX parsing using python-docx.

Word documents are structurally easier than PDFs — headings, lists,
and tables are already tagged in the XML (no font-size guessing needed).
The main extra work is:
  1. Walking the document body in DOCUMENT ORDER (paragraphs + tables
     are stored in separate collections by python-docx, so we must
     re-interleave them using the underlying XML).
  2. Extracting embedded images from the docx ZIP and auto-captioning.
"""

import os
import re
from typing import List
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from schema import DocElement, ElementType, BoundingBox, ParsedDocument
from captioner import caption_image


# ─────────────────────────────────────────────
# DOCUMENT-ORDER ITERATOR
# ─────────────────────────────────────────────
def iter_block_items(doc: Document):
    """
    Yield paragraphs and tables in the order they appear in the document.
    python-docx exposes doc.paragraphs and doc.tables separately —
    this walks the raw XML body so order is preserved.
    """
    parent_elm = doc.element.body
    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)


# ─────────────────────────────────────────────
# HEADING DETECTION (via Word style names)
# ─────────────────────────────────────────────
def get_heading_level(paragraph: Paragraph) -> int:
    style_name = (paragraph.style.name or "").lower()
    match = re.match(r"heading\s*(\d)", style_name)
    if match:
        return int(match.group(1))
    if style_name == "title":
        return 1
    return None


def is_list_item(paragraph: Paragraph) -> bool:
    style_name = (paragraph.style.name or "").lower()
    if "list" in style_name:
        return True
    # also check raw text bullet markers
    return bool(re.match(r"^\s*([\u2022\-\*]|\d+[\.\)])\s+", paragraph.text))


# ─────────────────────────────────────────────
# TABLE → MARKDOWN
# ─────────────────────────────────────────────
def docx_table_to_markdown(table: Table) -> str:
    rows = []
    for row in table.rows:
        rows.append([cell.text.strip().replace("\n", " ") for cell in row.cells])
    if not rows:
        return ""
    header = rows[0]
    md = "| " + " | ".join(header) + " |\n"
    md += "| " + " | ".join(["---"] * len(header)) + " |\n"
    for row in rows[1:]:
        md += "| " + " | ".join(row) + " |\n"
    return md.strip()


# ─────────────────────────────────────────────
# IMAGE EXTRACTION FROM DOCX
# ─────────────────────────────────────────────
def extract_docx_images(docx_path: str, doc_id: str, output_dir: str) -> dict:
    """
    DOCX images live inside the .docx ZIP under word/media/.
    Returns {rel_id: image_path} so we can later map inline images
    found in paragraph runs back to their saved files.
    """
    import zipfile
    os.makedirs(output_dir, exist_ok=True)
    image_map = {}

    with zipfile.ZipFile(docx_path) as z:
        media_files = [n for n in z.namelist() if n.startswith("word/media/")]
        for i, media_path in enumerate(media_files):
            ext = media_path.split(".")[-1]
            out_name = f"{doc_id}_img{i}.{ext}"
            out_path = os.path.join(output_dir, out_name)
            with open(out_path, "wb") as f:
                f.write(z.read(media_path))
            image_map[media_path] = out_path

    return image_map


# ─────────────────────────────────────────────
# MAIN PARSE FUNCTION
# ─────────────────────────────────────────────
def parse_docx(docx_path: str, output_dir: str = "./output/images") -> ParsedDocument:
    doc_id = os.path.splitext(os.path.basename(docx_path))[0]
    doc = Document(docx_path)

    print(f"[*] Parsing DOCX: {docx_path}")

    elements: List[DocElement] = []
    elem_counter = 0
    current_page = 1   # DOCX has no real page concept until rendered; we track section breaks as proxy

    # Extract all images first
    image_map = extract_docx_images(docx_path, doc_id, output_dir)
    image_paths = list(image_map.values())
    img_pointer = 0

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()

            # Check for inline images in this paragraph's runs
            has_image = bool(block._element.findall(".//" + qn("w:drawing")))

            if has_image and img_pointer < len(image_paths):
                elem_counter += 1
                img_path = image_paths[img_pointer]
                img_pointer += 1
                caption, conf = caption_image(img_path)
                elements.append(DocElement(
                    element_id=f"{doc_id}_el{elem_counter}_img",
                    type=ElementType.IMAGE,
                    content=caption,
                    page=current_page,
                    confidence=conf,
                    metadata={"image_path": img_path, "auto_captioned": True},
                ))

            if not text:
                continue

            elem_counter += 1
            heading_level = get_heading_level(block)

            if heading_level:
                etype = ElementType.HEADING
            elif is_list_item(block):
                etype = ElementType.LIST_ITEM
            elif re.match(r"^(Figure|Fig\.?|Table)\s+\d+[:.]\s*", text, re.IGNORECASE):
                etype = ElementType.CAPTION
            else:
                etype = ElementType.PARAGRAPH

            elements.append(DocElement(
                element_id=f"{doc_id}_el{elem_counter}",
                type=etype,
                content=text,
                page=current_page,
                heading_level=heading_level,
                metadata={"style": block.style.name},
            ))

            # Word inserts explicit page breaks as run properties — bump our proxy counter
            if block._element.findall(".//" + qn("w:br")):
                for br in block._element.findall(".//" + qn("w:br")):
                    if br.get(qn("w:type")) == "page":
                        current_page += 1

        elif isinstance(block, Table):
            elem_counter += 1
            md_table = docx_table_to_markdown(block)
            elements.append(DocElement(
                element_id=f"{doc_id}_el{elem_counter}_table",
                type=ElementType.TABLE,
                content=md_table,
                page=current_page,
                metadata={
                    "rows": len(block.rows),
                    "cols": len(block.columns),
                    "extraction_method": "python-docx",
                },
            ))

    # --- HIERARCHY TAGGING ---
    current_parents = {}
    for el in elements:
        if el.type == ElementType.HEADING:
            lvl = el.heading_level or 1
            current_parents[lvl] = el.element_id
            for l in range(lvl + 1, 7): current_parents.pop(l, None)
            for l in range(lvl - 1, 0, -1):
                if l in current_parents:
                    el.parent_id = current_parents[l]
                    break
        else:
            if current_parents:
                el.parent_id = current_parents[max(current_parents.keys())]
    # -------------------------

    parsed = ParsedDocument(
        source_file=docx_path,
        doc_type="docx",
        total_pages=current_page,
        elements=elements,
        parse_stats={"total_elements": len(elements), "by_type": {}},
    )
    parsed.parse_stats["by_type"] = parsed.summary()
    return parsed


if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Usage: python docx_parser.py <path_to_docx>")
        sys.exit(1)

    result = parse_docx(sys.argv[1])
    print(f"\n[✓] Parsed {result.parse_stats['total_elements']} elements")
    print(f"[✓] Breakdown: {result.parse_stats['by_type']}")
    result.save("./output/parsed_output_docx.json")
